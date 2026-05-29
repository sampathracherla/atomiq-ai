"""Generate Brownbag Session Word Document for Atomiq AI — Agentic Architecture Edition."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===== TITLE PAGE =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Atomiq AI')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Agentic Test Automation Framework')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Multi-Agent Orchestration \u2014 One Framework, All Platforms, AI-Powered')
run.italic = True
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('Brownbag Session Guide\n')
run.font.size = Pt(12)
run = author.add_run('By Sampath Racherla\n')
run.font.size = Pt(14)
run.bold = True
run = author.add_run('\ngithub.com/sampathracherla/atomiq-ai')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_page_break()

# ===== SECTION 1: WHAT IS ATOMIQ AI? =====
doc.add_heading('What is Atomiq AI?', level=1)
doc.add_paragraph(
    'Atomiq AI is an enterprise test automation framework that uses AI agents to autonomously '
    'run, manage, and report on tests across all platforms \u2014 Web, API, Mobile, and SAP.'
)
doc.add_paragraph()
doc.add_paragraph(
    'Think of it like a team of specialist robots: one knows how to test websites, one handles '
    'API endpoints, one tests on phones, and one deals with SAP. A supervisor robot coordinates '
    'them all and delivers a unified report.'
)

doc.add_heading('The Problem We Solve', level=2)
doc.add_paragraph('3-5 separate tools to license, learn, and maintain', style='List Bullet')
doc.add_paragraph('Selectors break every sprint \u2014 hours spent on manual fixes', style='List Bullet')
doc.add_paragraph('No AI integration \u2014 manual test creation and failure diagnosis', style='List Bullet')
doc.add_paragraph('Team silos \u2014 different skills needed per tool', style='List Bullet')
doc.add_paragraph('SAP teams use completely different tooling than web teams', style='List Bullet')

doc.add_page_break()

# ===== SECTION 2: AGENTIC ARCHITECTURE =====
doc.add_heading('Agentic Architecture', level=1)
doc.add_paragraph(
    'Atomiq AI uses a Supervisor pattern where a central orchestrator agent delegates work '
    'to specialist agents. Each agent is autonomous, communicates via a MessageBus, and '
    'reports results back to the supervisor.'
)

doc.add_heading('The Agents', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Agent'
table.rows[0].cells[1].text = 'Role'
table.rows[0].cells[2].text = 'Capabilities'
table.rows[1].cells[0].text = 'Supervisor'
table.rows[1].cells[1].text = 'Orchestrator'
table.rows[1].cells[2].text = 'orchestrate, run-regression, system-status'
table.rows[2].cells[0].text = 'Web Agent'
table.rows[2].cells[1].text = 'Browser testing'
table.rows[2].cells[2].text = 'run-all, run-spec, run-grep, list-specs'
table.rows[3].cells[0].text = 'API Agent'
table.rows[3].cells[1].text = 'REST/GraphQL'
table.rows[3].cells[2].text = 'run-all, run-spec, run-grep'
table.rows[4].cells[0].text = 'Mobile Agent'
table.rows[4].cells[1].text = 'Device testing'
table.rows[4].cells[2].text = 'run-all, run-spec, run-device'
table.rows[5].cells[0].text = 'SAP Agent'
table.rows[5].cells[1].text = 'Fiori/UI5'
table.rows[5].cells[2].text = 'run-all, run-spec, run-transaction, check-connection'

doc.add_heading('How They Communicate', level=2)
doc.add_paragraph(
    'All agents communicate through a MessageBus using two patterns:'
)
doc.add_paragraph('Pub/Sub \u2014 agents subscribe to their role and receive messages directed at them', style='List Bullet')
doc.add_paragraph('Request/Reply \u2014 supervisor sends a task and awaits the result with a correlation ID', style='List Bullet')
doc.add_paragraph('Broadcast \u2014 system-wide events visible to all agents', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    'Flow: Supervisor \u2192 MessageBus \u2192 WebAgent \u2192 executes tests \u2192 MessageBus \u2192 Supervisor receives result'
)

doc.add_page_break()

# ===== SECTION 3: FULL ARCHITECTURE STACK =====
doc.add_heading('Full Architecture Stack', level=1)
arch_text = (
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502  Agentic Orchestration Layer                             \u2502\n'
    '\u2502  Supervisor \u2192 Web | API | Mobile | SAP Agents            \u2502\n'
    '\u2502  MessageBus (pub/sub + request/reply) + AgentRegistry    \u2502\n'
    '\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524\n'
    '\u2502  Self-Healing Engine | Visual Regression | Data Generator \u2502\n'
    '\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524\n'
    '\u2502  AI Engine (OpenAI / Azure / Gemini / Claude / Custom)   \u2502\n'
    '\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524\n'
    '\u2502  Playwright Foundation (Web, API, Mobile, SAP adapters)  \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518'
)
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Technology Stack', level=2)
doc.add_paragraph('TypeScript (strict mode, ES2022)', style='List Bullet')
doc.add_paragraph('Playwright (Microsoft browser automation)', style='List Bullet')
doc.add_paragraph('5 AI Providers (OpenAI, Azure OpenAI, Gemini, Claude, Custom)', style='List Bullet')
doc.add_paragraph('Node.js runtime', style='List Bullet')
doc.add_paragraph('Open source \u2014 $0 licensing', style='List Bullet')

doc.add_page_break()

# ===== SECTION 4: KEY FEATURES =====
doc.add_heading('Key Features', level=1)

doc.add_heading('1. Multi-Agent Orchestration', level=2)
doc.add_paragraph(
    'A supervisor agent coordinates specialist agents. Run all tests with one command, '
    'or target a single agent. Agents are autonomous \u2014 they execute, report metrics, '
    'and can be extended with new capabilities.'
)

doc.add_heading('2. Self-Healing Selectors', level=2)
doc.add_paragraph(
    'When a selector breaks (developer renamed a button), the framework tries 4 strategies: '
    'attribute fallback, text content, structural similarity, and AI-assisted healing. '
    'Tests keep running without human intervention.'
)

doc.add_heading('3. Pluggable AI Providers', level=2)
doc.add_paragraph(
    'Switch between OpenAI, Azure OpenAI, Google Gemini, and Anthropic Claude with one config '
    'change. AI generates tests from descriptions, diagnoses failures, and suggests fixes. '
    'No vendor lock-in.'
)

doc.add_heading('4. Page Object Model (POM)', level=2)
doc.add_paragraph(
    'BasePage class provides self-healing, visual regression, and element discovery. '
    'Locators defined once, reused across all tests. Change one file, not 100 tests.'
)

doc.add_heading('5. Visual Regression Testing', level=2)
doc.add_paragraph(
    'Pixel-level screenshot comparison with configurable thresholds. '
    'Auto-baseline management and the ability to mask dynamic regions.'
)

doc.add_page_break()

# ===== SECTION 5: LIVE DEMO =====
doc.add_heading('Live Demo', level=1)

doc.add_heading('Run All Agents', level=2)
doc.add_paragraph('Command: npx tsx examples/agent-demo.ts')
doc.add_paragraph()
doc.add_paragraph('Expected output:')
doc.add_paragraph('  WebAgent:    \u2705 PASSED (5/5 passed in 9.3s)', style='List Bullet')
doc.add_paragraph('  ApiAgent:    \u2705 PASSED (5/5 passed in 8.2s)', style='List Bullet')
doc.add_paragraph('  MobileAgent: \u2705 PASSED (3/3 passed in 7.1s)', style='List Bullet')
doc.add_paragraph('  SapAgent:    \u26a0\ufe0f  Not configured (no SAP env)', style='List Bullet')
doc.add_paragraph('  Total: 13 tests, 100% pass rate, ~24.5s', style='List Bullet')

doc.add_heading('Run Single Agent', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Command'
table.rows[0].cells[1].text = 'What It Runs'
table.rows[1].cells[0].text = 'npx tsx examples/agent-demo.ts web'
table.rows[1].cells[1].text = 'SauceDemo e-commerce tests (5 tests)'
table.rows[2].cells[0].text = 'npx tsx examples/agent-demo.ts api'
table.rows[2].cells[1].text = 'REST API CRUD tests (5 tests)'
table.rows[3].cells[0].text = 'npx tsx examples/agent-demo.ts mobile'
table.rows[3].cells[1].text = 'Device viewport tests (3 tests)'
table.rows[4].cells[0].text = 'npx tsx examples/agent-demo.ts sap'
table.rows[4].cells[1].text = 'SAP connectivity check'

doc.add_heading('Run Tests Directly (no agent layer)', level=2)
doc.add_paragraph('npx playwright test examples/saucedemo-test.spec.ts', style='List Bullet')
doc.add_paragraph('npx playwright test examples/api-test.spec.ts', style='List Bullet')
doc.add_paragraph('npx playwright test examples/mobile-test.spec.ts', style='List Bullet')
doc.add_paragraph('npx playwright show-report', style='List Bullet')

doc.add_page_break()

# ===== SECTION 6: RESULTS =====
doc.add_heading('Results at a Glance', level=1)

table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Metric'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'Total Tests Passing'
table.rows[1].cells[1].text = '13 (Web 5, API 5, Mobile 3)'
table.rows[2].cells[0].text = 'Specialist Agents'
table.rows[2].cells[1].text = '4 (Web, API, Mobile, SAP)'
table.rows[3].cells[0].text = 'Pass Rate'
table.rows[3].cells[1].text = '100%'
table.rows[4].cells[0].text = 'Total Runtime'
table.rows[4].cells[1].text = '~24.5 seconds'
table.rows[5].cells[0].text = 'AI Providers Supported'
table.rows[5].cells[1].text = '5'
table.rows[6].cells[0].text = 'Licensing Cost'
table.rows[6].cells[1].text = '$0 (open source)'

doc.add_page_break()

# ===== SECTION 7: BUSINESS VALUE =====
doc.add_heading('Business Value & ROI', level=1)
doc.add_paragraph('Eliminate tool sprawl \u2014 replaces 3-5 separate tools with one framework', style='List Bullet')
doc.add_paragraph('80% less maintenance \u2014 self-healing eliminates manual selector fixes', style='List Bullet')
doc.add_paragraph('5x faster test creation \u2014 AI generates tests from natural language', style='List Bullet')
doc.add_paragraph('One skill set \u2014 TypeScript covers all platforms, no team silos', style='List Bullet')
doc.add_paragraph('Autonomous operation \u2014 agents run, diagnose, and report without intervention', style='List Bullet')
doc.add_paragraph('Faster releases \u2014 reliable tests = confidence to ship frequently', style='List Bullet')
doc.add_paragraph('Enterprise-ready \u2014 SAP Fiori, REST APIs, responsive mobile, visual regression', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    'Bottom line: Faster releases, fewer flaky tests, lower licensing costs, '
    'one team can test everything.'
)
run.bold = True

doc.add_page_break()

# ===== SECTION 8: ROADMAP =====
doc.add_heading('Roadmap', level=1)

doc.add_heading('Completed', level=2)
doc.add_paragraph('Phase 1: Core foundation \u2014 BaseAgent, MessageBus, AgentRegistry, Types', style='List Bullet')
doc.add_paragraph('Phase 2: Supervisor + WebAgent \u2014 orchestration + browser testing', style='List Bullet')
doc.add_paragraph('Phase 3: API Agent, Mobile Agent, SAP Agent \u2014 full specialist coverage', style='List Bullet')

doc.add_heading('Coming Next', level=2)
doc.add_paragraph('Phase 4: Planner Agent \u2014 AI-powered test planning from requirements', style='List Bullet')
doc.add_paragraph('Phase 4: Healer Agent \u2014 auto-diagnose and fix failing tests', style='List Bullet')
doc.add_paragraph('Phase 5: Report Agent \u2014 AI-generated analytics and recommendations', style='List Bullet')
doc.add_paragraph('CI/CD integration \u2014 GitHub Actions, Azure DevOps', style='List Bullet')
doc.add_paragraph('SAP system pilot \u2014 connect to real SAP environment', style='List Bullet')

doc.add_page_break()

# ===== SECTION 9: Q&A =====
doc.add_heading('Anticipated Questions', level=1)

qa_items = [
    ('How is this different from Selenium?',
     'Selenium is 20 years old, no AI, no self-healing, no agent orchestration. Playwright (our foundation) is Microsoft\'s modern replacement with multi-browser support.'),
    ('What makes the "agentic" approach special?',
     'Each agent is autonomous and specialized. The supervisor coordinates them like a team lead \u2014 you get parallel execution, automatic routing, and a single entry point for all testing.'),
    ('Does it work with SAP?',
     'Yes, there\'s a dedicated SAP Agent for Fiori/UI5 apps. It uses dhikraft patterns for UI5 elements.'),
    ('What skills does my team need?',
     'TypeScript basics. One language covers all platforms \u2014 no Java for Selenium, no Python for API, no separate SAP tool.'),
    ('Can it run in CI/CD?',
     'Yes \u2014 GitHub Actions, Azure DevOps, Jenkins \u2014 all supported natively by Playwright.'),
    ('What about test data security?',
     'Credentials stored in environment variables, never in code. Data generation creates synthetic data, never uses real PII.'),
]

table = doc.add_table(rows=len(qa_items) + 1, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Question'
table.rows[0].cells[1].text = 'Answer'
for i, (q, a) in enumerate(qa_items, 1):
    table.rows[i].cells[0].text = q
    table.rows[i].cells[1].text = a

doc.add_page_break()

# ===== ELEVATOR PITCH =====
doc.add_heading('30-Second Elevator Pitch', level=1)
doc.add_paragraph()
pitch = doc.add_paragraph()
pitch.paragraph_format.left_indent = Cm(1)
pitch.paragraph_format.right_indent = Cm(1)
run = pitch.add_run(
    '"Atomiq AI is an agentic test automation framework. It has AI agents that autonomously '
    'test your web apps, APIs, mobile layouts, and SAP systems. A supervisor agent orchestrates '
    'everything \u2014 run one command and get a full regression across all platforms in under 30 seconds. '
    'It self-heals when selectors break, uses AI to generate tests from English descriptions, '
    'and costs nothing to license. I built it."'
)
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('GitHub: github.com/sampathracherla/atomiq-ai')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

# Save
import os
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'BROWNBAG-SESSION-GUIDE.docx')
doc.save(output_path)
print(f'Word document saved: {output_path}')
